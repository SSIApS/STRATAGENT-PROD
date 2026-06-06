"""
STRATAGENT -- Word Document Export Service
Generates a branded .docx from FI output results (Value Brief, First Signal, Convergence Proposal).
"""
import io
import pathlib
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# SSI brand colours
GOLD   = RGBColor(0xC9, 0xA2, 0x27)
DARK   = RGBColor(0x1A, 0x1A, 0x2E)
MUTED  = RGBColor(0x64, 0x74, 0x8B)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)

LOGO_PATH = pathlib.Path(__file__).parent.parent.parent / "frontend" / "public" / "stratagent-logo.png"


def _add_horizontal_rule(doc: Document):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "C9A227")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def _set_font(run, size_pt=11, bold=False, color=None, italic=False):
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = "Arial"
    if color:
        run.font.color.rgb = color


def generate_output_brief(
    output_data: dict,
    company_name: str,
    supplier_name: str,
    convergence_index: int,
    path_label: str,
) -> bytes:
    """
    Generate a Word doc from a STRATAGENT output result.
    Returns the .docx as bytes for HTTP response.
    """
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # ── LOGO ─────────────────────────────────────────────────────────────
    if LOGO_PATH.exists():
        logo_para = doc.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = logo_para.add_run()
        run.add_picture(str(LOGO_PATH), width=Inches(2.8))
    else:
        p = doc.add_paragraph("STRATAGENT")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_font(p.runs[0], size_pt=22, bold=True, color=GOLD)

    # ── HEADER INFO ──────────────────────────────────────────────────────
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run(path_label)
    _set_font(run, size_pt=18, bold=True, color=DARK)

    meta = doc.add_paragraph()
    run = meta.add_run(
        f"Prospect: {company_name}  |  Supplier: {supplier_name}  |  "
        f"Convergence Index: {convergence_index}  |  {datetime.now().strftime('%d %B %Y')}"
    )
    _set_font(run, size_pt=10, color=MUTED)

    _add_horizontal_rule(doc)
    doc.add_paragraph()

    # ── OUTREACH EMAIL ───────────────────────────────────────────────────
    email_text = output_data.get("email", "")
    if email_text:
        h = doc.add_paragraph()
        run = h.add_run("OUTREACH EMAIL")
        _set_font(run, size_pt=11, bold=True, color=GOLD)

        note = doc.add_paragraph()
        run = note.add_run("Ready to send. Copy into your email client.")
        _set_font(run, size_pt=9, italic=True, color=MUTED)

        doc.add_paragraph()
        email_box = doc.add_paragraph()
        email_box.paragraph_format.left_indent  = Inches(0.3)
        email_box.paragraph_format.right_indent = Inches(0.3)

        # Strip the SSI footer from the email body for cleaner display
        # (footer is already in the doc footer)
        email_clean = email_text.split("---\nJason")[0].strip()
        for line in email_clean.split("\n"):
            run = email_box.add_run(line + "\n")
            _set_font(run, size_pt=10.5)

        doc.add_paragraph()
        _add_horizontal_rule(doc)
        doc.add_paragraph()

    # ── VALUE BRIEF ──────────────────────────────────────────────────────
    brief_text = output_data.get("brief", "")
    if brief_text:
        h = doc.add_paragraph()
        run = h.add_run("MUTUAL VALUE BRIEF")
        _set_font(run, size_pt=11, bold=True, color=GOLD)
        doc.add_paragraph()

        brief_clean = brief_text.split("---\nJason")[0].strip()
        for line in brief_clean.split("\n"):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(line)
            _set_font(run, size_pt=10.5)

        doc.add_paragraph()
        _add_horizontal_rule(doc)
        doc.add_paragraph()

    # ── PROPOSAL ─────────────────────────────────────────────────────────
    proposal_text = output_data.get("proposal", "")
    if proposal_text:
        h = doc.add_paragraph()
        run = h.add_run("TECHNICAL PROPOSAL")
        _set_font(run, size_pt=11, bold=True, color=GOLD)
        doc.add_paragraph()

        proposal_clean = proposal_text.split("---\nJason")[0].strip()
        for line in proposal_clean.split("\n"):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(line)
            _set_font(run, size_pt=10.5)

        doc.add_paragraph()
        _add_horizontal_rule(doc)
        doc.add_paragraph()

    # ── ENGAGEMENT BRIEF / RFQ ───────────────────────────────────────────
    rfq_text = output_data.get("engagement_brief", "")
    if rfq_text:
        h = doc.add_paragraph()
        run = h.add_run("ENGAGEMENT BRIEF / RFQ FRAMEWORK")
        _set_font(run, size_pt=11, bold=True, color=GOLD)
        doc.add_paragraph()

        rfq_clean = rfq_text.split("---\nJason")[0].strip()
        for line in rfq_clean.split("\n"):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(line)
            _set_font(run, size_pt=10.5)

        doc.add_paragraph()
        _add_horizontal_rule(doc)
        doc.add_paragraph()

    # ── QUALIFYING QUESTIONS ─────────────────────────────────────────────
    questions = output_data.get("qualifying_questions", [])
    if questions:
        h = doc.add_paragraph()
        run = h.add_run("QUALIFYING QUESTIONS  —  First Call Preparation")
        _set_font(run, size_pt=11, bold=True, color=GOLD)

        sub = doc.add_paragraph()
        run = sub.add_run("Use these to uncover the specific need and qualify the opportunity.")
        _set_font(run, size_pt=9, italic=True, color=MUTED)
        doc.add_paragraph()

        for i, q in enumerate(questions, 1):
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(q)
            _set_font(run, size_pt=10.5)

        doc.add_paragraph()

    # ── SSI FOOTER ───────────────────────────────────────────────────────
    _add_horizontal_rule(doc)
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run(
        "Jason L. Smith  |  Strategic Sales International ApS  |  "
        "info@strategic.dk  |  +45 24 99 23 93  |  www.strategic-dk.com\n"
        "CVR: 41945621  |  Roskilde, Denmark  |  Generated by STRATAGENT"
    )
    _set_font(run, size_pt=8.5, color=MUTED)

    # ── RETURN AS BYTES ──────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
