"""
STRATAGENT -- Output Engine Router
Graduated document generation based on Singularity Density (SD) score.
Path A: CONVERGENCE PROPOSAL (90-100)
Path B: MUTUAL VALUE BRIEF (75-89)
Path C: FIRST SIGNAL (60-74)
"""
import io
import os
import re
from datetime import date
from typing import Any, Dict

from fastapi import APIRouter, HTTPException, Header
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

from services import firestore as db
from services.demo_gate import check_and_increment
from agents.output_agent import (
    generate_convergence_proposal,
    generate_mutual_value_brief,
    generate_first_signal,
)

router = APIRouter()

_LOGO_PATH = os.path.normpath(os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    "..", "..", "..", "Marketing", "Logos", "Stratagent_Orange_Standard.png"
))

ORANGE     = None   # set lazily inside functions after docx import
DARK_GREY  = None
LIGHT_GREY = None

# ---------------------------------------------------------------------------
# Text cleaning
# ---------------------------------------------------------------------------

_SSI_FOOTER_RE = re.compile(
    r'\n{0,3}(?:-{3,}\n+)?Jason L\. Smith[^\n]*\n(?:[^\n]+\n){0,5}[^\n]*STRATAGENT[^\n]*\.?',
    re.IGNORECASE,
)

def _clean(text: str) -> str:
    """Strip Gemini auto-footers and known placeholder artifacts."""
    if not text:
        return text
    # Remove all SSI footer blocks
    text = _SSI_FOOTER_RE.sub('', text)
    # Remove dangling dividers
    text = re.sub(r'\n+---+\s*$', '', text)
    # Replace sender placeholder
    text = re.sub(r'\[Your Name[^\]]*\]', 'Jason L. Smith | SSI ApS', text)
    # Replace date placeholders with actual date
    for ph in ('[Current Date]', '[Date]', '[date]'):
        text = text.replace(ph, date.today().strftime('%d %B %Y'))
    return text.strip()


# ---------------------------------------------------------------------------
# Markdown → docx renderer
# ---------------------------------------------------------------------------

def _render_markdown(doc, text: str):
    """Parse lightweight markdown → properly formatted Word content."""
    from docx.shared import Pt, RGBColor, Inches

    ORANGE_C = RGBColor(0xE8, 0x7A, 0x00)

    def parse_inline(para, content: str):
        parts = re.split(r'(\*\*[^*\n]+\*\*)', content)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = para.add_run(part[2:-2])
                run.bold = True
            elif part:
                para.add_run(part)

    for line in text.splitlines():
        stripped = line.strip()

        if not stripped or re.match(r'^[-*_]{3,}$', stripped):
            continue

        if stripped.startswith('## '):
            p = doc.add_paragraph()
            r = p.add_run(stripped[3:].strip())
            r.bold = True; r.font.size = Pt(13); r.font.color.rgb = ORANGE_C
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)

        elif stripped.startswith('### '):
            p = doc.add_paragraph()
            r = p.add_run(stripped[4:].strip())
            r.bold = True; r.font.size = Pt(11); r.font.color.rgb = ORANGE_C
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(3)

        elif stripped.startswith('#### '):
            p = doc.add_paragraph()
            r = p.add_run(stripped[5:].strip())
            r.bold = True; r.font.size = Pt(10); r.font.color.rgb = ORANGE_C
            p.paragraph_format.space_before = Pt(6)

        elif re.match(r'^\s{0,8}[\*\-]\s', line):
            depth = 1 if (len(line) - len(line.lstrip())) >= 4 else 0
            content = re.sub(r'^\s*[\*\-]\s+', '', line)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3 + depth * 0.22)
            p.paragraph_format.first_line_indent = Inches(-0.18)
            p.paragraph_format.space_after = Pt(2)
            p.add_run('• ')
            parse_inline(p, content)

        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(5)
            parse_inline(p, stripped)


# ---------------------------------------------------------------------------
# Docx helpers
# ---------------------------------------------------------------------------

def _shade_cell(cell, hex_color: str):
    """Apply background fill to a table cell via OOXML."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _cell_text(cell, text: str, bold=False, color=None, size_pt=10, align='left'):
    """Set text in a table cell with formatting."""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    align_map = {
        'left': WD_ALIGN_PARAGRAPH.LEFT,
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
    }
    p.alignment = align_map.get(align, WD_ALIGN_PARAGRAPH.LEFT)


def _sd_bar(score: int, total_blocks: int = 20) -> str:
    """Return a Unicode block progress bar for the SD score."""
    filled = round(score / 100 * total_blocks)
    return '█' * filled + '░' * (total_blocks - filled)


def _sd_label(score: int) -> str:
    if score >= 90: return 'CONVERGENCE READY'
    if score >= 75: return 'VALUE BRIEF READY'
    if score >= 60: return 'FIRST SIGNAL'
    return 'PARKED'


# ---------------------------------------------------------------------------
# Main docx builder
# ---------------------------------------------------------------------------

def _signal_strength_color(strength: str) -> str:
    """Map a buying-signal strength to a hex accent color for the docx."""
    s = (strength or '').upper()
    if s == 'HIGH':
        return 'B23B00'
    if s == 'MEDIUM':
        return 'E87A00'
    return '9A9A9A'


# ---------------------------------------------------------------------------
# Brand system -- shared across ALL STRATAGENT docx report types
# ---------------------------------------------------------------------------

# A4 content width at 0.75in margins: 8.27 - 1.5 = 6.77in
_CONTENT_W = 6.77


def _page_setup(doc) -> None:
    """A4, 0.75in margins -- consistent across all report types."""
    from docx.shared import Mm, Inches
    for sec in doc.sections:
        sec.page_width    = Mm(210)
        sec.page_height   = Mm(297)
        sec.top_margin    = Inches(0.5)
        sec.bottom_margin = Inches(0.5)
        sec.left_margin   = Inches(0.5)
        sec.right_margin  = Inches(0.5)


def _set_col_widths(tbl, widths_inches: list) -> None:
    """Set explicit column widths (clips silently if more widths than columns)."""
    from docx.shared import Inches
    for i, col in enumerate(tbl.columns):
        if i < len(widths_inches):
            col.width = Inches(widths_inches[i])


def _cell_margins(cell, top_pt: float = 4, right_pt: float = 6,
                  bot_pt: float = 4, left_pt: float = 6) -> None:
    """Set inner cell padding in points."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcMar')):
        tcPr.remove(old)
    tcMar = OxmlElement('w:tcMar')
    for side, pt in [('top', top_pt), ('right', right_pt),
                     ('bottom', bot_pt), ('left', left_pt)]:
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:w'), str(int(pt * 20)))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def _keep_with_next_para(p) -> None:
    """Prevent a paragraph from being separated from the one following it."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = p._p.get_or_add_pPr()
    kwn = OxmlElement('w:keepWithNext')
    pPr.append(kwn)


def _no_table_borders(tbl) -> None:
    """Remove all outer and inner borders from a table."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tblPr = tbl._tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl._tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'none')
        b.set(qn('w:sz'), '0')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'auto')
        tblBorders.append(b)
    for existing in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(existing)
    tblPr.append(tblBorders)


def _add_page_footer(doc) -> None:
    """No-op -- footer appended as final paragraph by each builder via _doc_footer()."""
    pass


def _doc_footer(doc) -> None:
    """SSI branding text block appended at the end of every report."""
    from docx.shared import Pt, RGBColor
    doc.add_paragraph()
    hr = doc.add_paragraph()
    r = hr.add_run('\u2500' * 72)
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    ft = doc.add_paragraph()
    r = ft.add_run(
        'Jason L. Smith  |  Strategic Sales International ApS\n'
        'info@strategic.dk  |  www.strategic-dk.com  |  +45 24 99 23 93\n'
        'CVR: 41945621  |  Roskilde, Denmark\n'
        'STRATAGENT -- The Intelligence Behind Agentic Sales'
    )
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x9A, 0x9A, 0x9A)
def _cover_block(doc, report_type: str, company: str,
                 location: str = "", subtitle: str = "") -> None:
    """Light cover -- matches STRATAGENT FI report style, printer-friendly."""
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    ORANGE_C = RGBColor(0xE8, 0x7A, 0x00)
    DARK_C   = RGBColor(0x1A, 0x1A, 0x1A)
    GRAY_C   = RGBColor(0x64, 0x74, 0x8B)
    LTGRAY_C = RGBColor(0x94, 0xA3, 0xB8)
    RULE_C   = RGBColor(0xCC, 0xCC, 0xCC)

    # Logo
    if os.path.exists(_LOGO_PATH):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.add_run().add_picture(_LOGO_PATH, width=Inches(2.2))
    else:
        p = doc.add_paragraph()
        r = p.add_run('STRATAGENT')
        r.bold = True; r.font.size = Pt(18); r.font.color.rgb = ORANGE_C

    doc.add_paragraph()

    # Report type title
    h = doc.add_paragraph()
    r = h.add_run(report_type)
    r.bold = True; r.font.size = Pt(22); r.font.color.rgb = ORANGE_C

    # Company name
    co = doc.add_paragraph()
    r = co.add_run(company)
    r.bold = True; r.font.size = Pt(14); r.font.color.rgb = DARK_C

    if location:
        lp = doc.add_paragraph()
        r = lp.add_run(location)
        r.font.size = Pt(10); r.font.color.rgb = GRAY_C

    if subtitle:
        sp = doc.add_paragraph()
        r = sp.add_run(subtitle)
        r.font.size = Pt(10); r.font.color.rgb = GRAY_C

    # Date / SSI line
    dl = doc.add_paragraph()
    r = dl.add_run(
        f"Prepared: {date.today().strftime('%d %B %Y')}"
        f"  ·  Strategic Sales International ApS"
    )
    r.font.size = Pt(8.5); r.font.color.rgb = LTGRAY_C

    # Horizontal rule
    hr = doc.add_paragraph()
    r = hr.add_run('─' * 72)
    r.font.size = Pt(8); r.font.color.rgb = RULE_C

    doc.add_paragraph()


def _section_heading(doc, text: str) -> None:
    """Orange section bar -- white text, printer-friendly, prevents page splits."""
    from docx.shared import Pt, RGBColor

    WHITE_C = RGBColor(0xFF, 0xFF, 0xFF)

    # Pre-spacer with keep_with_next so heading never orphans at page bottom
    pre = doc.add_paragraph()
    pre.paragraph_format.space_before = Pt(0)
    pre.paragraph_format.space_after  = Pt(0)
    _keep_with_next_para(pre)

    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    _shade_cell(cell, 'E87A00')

    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    _keep_with_next_para(p)
    r = p.add_run(text.upper())
    r.bold = True; r.font.size = Pt(10)
    r.font.color.rgb = WHITE_C

    # Post-spacer
    post = doc.add_paragraph()
    post.paragraph_format.space_before = Pt(0)
    post.paragraph_format.space_after  = Pt(4)


def _body_para(doc, text: str, italic: bool = False,
               size: float = 10.5, color: str = '1E293B'):
    """Standard body paragraph."""
    from docx.shared import Pt, RGBColor
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return p


def _score_table(doc, rows: list, col_widths: list = None) -> None:
    """Score table: label | bar | value. rows=[(label, val, max_val), ...]"""
    from docx.shared import Pt, RGBColor
    if not rows:
        return

    cw = col_widths or [2.4, 3.6, 0.77]
    tbl = doc.add_table(rows=len(rows), cols=3)
    tbl.style = 'Table Grid'
    _set_col_widths(tbl, cw)

    for i, (label, val, max_val) in enumerate(rows):
        pct = int((val / max_val) * 100) if max_val else 0
        if pct >= 75:   col_hex = '10B981'
        elif pct >= 50: col_hex = 'F59E0B'
        else:           col_hex = 'EF4444'

        bg = 'FFFFFF' if i % 2 == 0 else 'F8FAFC'
        cells = tbl.rows[i].cells

        _shade_cell(cells[0], bg)
        _cell_margins(cells[0], 4, 8, 4, 8)
        _cell_text(cells[0], label, size_pt=9.5, color='1E293B')

        bar_len = max(1, int(pct / 5))
        _shade_cell(cells[1], bg)
        _cell_margins(cells[1], 6, 8, 6, 8)
        _cell_text(cells[1], '█' * bar_len, size_pt=7, color=col_hex)

        _shade_cell(cells[2], bg)
        _cell_margins(cells[2], 4, 8, 4, 8)
        _cell_text(cells[2], str(val), bold=True, size_pt=9.5,
                   color=col_hex, align='center')

    post = doc.add_paragraph()
    post.paragraph_format.space_after = Pt(4)


def _stats_table(doc, stats: list, col_widths: list = None) -> None:
    """Key/value stats table. stats=[(key, value), ...]"""
    from docx.shared import Pt
    if not stats:
        return

    cw = col_widths or [1.8, 4.97]
    tbl = doc.add_table(rows=len(stats), cols=2)
    tbl.style = 'Table Grid'
    _set_col_widths(tbl, cw)

    for i, (key, val) in enumerate(stats):
        _shade_cell(tbl.rows[i].cells[0], 'F1F5F9')
        _cell_margins(tbl.rows[i].cells[0], 4, 8, 4, 8)
        _cell_text(tbl.rows[i].cells[0], key, bold=True, size_pt=8.5, color='64748B')
        _cell_margins(tbl.rows[i].cells[1], 4, 8, 4, 8)
        _cell_text(tbl.rows[i].cells[1], str(val), size_pt=10)

    post = doc.add_paragraph()
    post.paragraph_format.space_after = Pt(4)


def _bullet_item(doc, text: str, bullet: str = '▸',
                 color: str = '1E293B') -> None:
    """Single formatted bullet point."""
    from docx.shared import Pt, RGBColor, Inches
    p = doc.add_paragraph()
    p.paragraph_format.left_indent        = Inches(0.2)
    p.paragraph_format.first_line_indent  = Inches(-0.2)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(f"{bullet}  {text}")
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor(*bytes.fromhex(color))


def _bullet_list(doc, items: list, bullet: str = '▸',
                 color: str = '1E293B') -> None:
    """Render a list of strings as formatted bullet points."""
    from docx.shared import Pt
    for item in items:
        _bullet_item(doc, str(item), bullet, color)
    post = doc.add_paragraph()
    post.paragraph_format.space_after = Pt(4)


def _channel_fit_table(doc, channel_fit: dict) -> None:
    """3-col channel fit: channel | fit | assessment."""
    from docx.shared import Pt, RGBColor
    if not channel_fit:
        return

    ch_labels = {
        "etsy":               "Etsy",
        "redbubble_society6": "Redbubble / Society6",
        "amazon":             "Amazon",
        "specialty_retail":   "Specialty Retail",
        "social_commerce":    "Social Commerce",
    }
    rows = [
        (ch_labels.get(k, k), v.get("fit", ""), v.get("reason", ""))
        for k, v in channel_fit.items() if isinstance(v, dict)
    ]
    if not rows:
        return

    cw = [1.5, 0.8, 4.47]
    tbl = doc.add_table(rows=len(rows) + 1, cols=3)
    tbl.style = 'Table Grid'
    _set_col_widths(tbl, cw)

    # Header row
    for cell, txt in zip(tbl.rows[0].cells, ['CHANNEL', 'FIT', 'ASSESSMENT']):
        _shade_cell(cell, 'E87A00')
        _cell_margins(cell, 5, 8, 5, 8)
        _cell_text(cell, txt, bold=True, size_pt=8.5, color='FFFFFF')

    fit_colors = {'high': '10B981', 'medium': 'F59E0B', 'low': 'EF4444'}
    for i, (ch_name, fit, reason) in enumerate(rows, start=1):
        bg = 'FFFFFF' if i % 2 == 1 else 'F8FAFC'
        cells = tbl.rows[i].cells
        _shade_cell(cells[0], bg)
        _cell_margins(cells[0], 4, 8, 4, 8)
        _cell_text(cells[0], ch_name, bold=True, size_pt=9.5)

        _shade_cell(cells[1], bg)
        _cell_margins(cells[1], 4, 8, 4, 8)
        fit_color = fit_colors.get((fit or '').lower(), '475569')
        _cell_text(cells[1], (fit or '').upper(), bold=True,
                   size_pt=9, color=fit_color, align='center')

        _shade_cell(cells[2], bg)
        _cell_margins(cells[2], 4, 8, 4, 8)
        _cell_text(cells[2], reason, size_pt=9.5)

    post = doc.add_paragraph()
    post.paragraph_format.space_after = Pt(4)


# ---------------------------------------------------------------------------
# FI output builder (Convergence Proposal / Mutual Value Brief / First Signal)
# ---------------------------------------------------------------------------

def _build_docx(label: str, company_name: str, supplier_name: str,
                convergence_index: int, output: dict,
                buying_signals=None, score_reasoning=None,
                what_would_improve=None) -> bytes:
    """Build a premium branded .docx from FI-generated output."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise RuntimeError("python-docx not installed.")

    ORANGE_C = RGBColor(0xE8, 0x7A, 0x00)
    WHITE_C  = RGBColor(0xFF, 0xFF, 0xFF)
    DARK_C   = RGBColor(0x1A, 0x1A, 0x1A)

    doc = Document()
    _page_setup(doc)
    _add_page_footer(doc)

    # Cover block
    _cover_block(doc, label, company_name)

    # Metadata block: supplier, SD score, date
    meta = [
        ('SUPPLIER',             supplier_name),
        ('SINGULARITY DENSITY',  f'{convergence_index}/100 -- {_sd_label(convergence_index)}'),
        ('DATE',                 date.today().strftime('%d %B %Y')),
        ('PREPARED BY',          'Jason L. Smith  |  Strategic Sales International ApS'),
    ]
    _stats_table(doc, meta, col_widths=[2.0, 4.77])

    # SD score bar
    bar_tbl = doc.add_table(rows=1, cols=1)
    bar_tbl.style = 'Table Grid'
    bar_cell = bar_tbl.rows[0].cells[0]
    _shade_cell(bar_cell, 'FFF3E0')
    bp = bar_cell.paragraphs[0]
    bp.clear()
    br = bp.add_run(_sd_bar(convergence_index))
    br.font.size = Pt(10)
    br.font.color.rgb = ORANGE_C

    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(10)

    # ── CONTENT SECTIONS ─────────────────────────────────────────────────────

    def add_section(title: str, body: str):
        _section_heading(doc, title)
        _render_markdown(doc, _clean(body))
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(6)

    if output.get('email'):
        add_section('Outreach Email', output['email'])

    if output.get('brief'):
        add_section('Value Brief', output['brief'])

    if output.get('proposal'):
        add_section('Technical Proposal', output['proposal'])

    if output.get('engagement_brief'):
        add_section('Engagement Brief / RFQ Framework', output['engagement_brief'])

    if output.get('qualifying_questions'):
        _section_heading(doc, 'Qualifying Questions')
        for i, q in enumerate(output['qualifying_questions'], 1):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.15)
            p.paragraph_format.space_after = Pt(5)
            num = p.add_run(f'{i}.  ')
            num.bold = True; num.font.color.rgb = ORANGE_C
            p.add_run(str(q)).font.size = Pt(10.5)
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    if buying_signals:
        _section_heading(doc, 'Buying Signals Detected')
        for sig in buying_signals:
            if not isinstance(sig, dict):
                continue
            sig_type = str(sig.get('type') or 'SIGNAL').upper()
            strength = str(sig.get('strength') or '').upper()
            description = sig.get('signal') or ''
            timing = sig.get('timing')
            source = sig.get('source')

            head_p = doc.add_paragraph()
            head_p.paragraph_format.space_before = Pt(6)
            head_p.paragraph_format.space_after  = Pt(1)
            tr = head_p.add_run(sig_type)
            tr.bold = True; tr.font.size = Pt(9)
            tr.font.color.rgb = ORANGE_C
            if strength:
                head_p.add_run('   ')
                sr = head_p.add_run(f'{strength} STRENGTH')
                sr.bold = True; sr.font.size = Pt(8)
                sr.font.color.rgb = RGBColor(*bytes.fromhex(_signal_strength_color(strength)))

            if description:
                _body_para(doc, str(description))

            extra = []
            if timing: extra.append(f'Timing: {timing}')
            if source:  extra.append(f'Source: {source}')
            if extra:
                mp = doc.add_paragraph()
                mp.paragraph_format.space_after = Pt(4)
                mr = mp.add_run('  |  '.join(extra))
                mr.italic = True; mr.font.size = Pt(8.5)
                mr.font.color.rgb = RGBColor(0x6B, 0x6B, 0x6B)

        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    if score_reasoning:
        _section_heading(doc, 'Singularity Density -- Score Reasoning')
        _body_para(doc, str(score_reasoning))
        if what_would_improve:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            r = p.add_run('What would raise this score:')
            r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = ORANGE_C
            _body_para(doc, str(what_would_improve))
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    _doc_footer(doc)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ---------------------------------------------------------------------------
# Routes
# ---------------------------------------------------------------------------

class GenerateRequest(BaseModel):
    profile_id: str

class ExportRequest(BaseModel):
    profile_id: str
    label: str
    company_name: str
    convergence_index: int
    supplier_name: str
    output: Dict[str, Any]
    include_buying_signals: bool = False
    include_reasoning: bool = False


@router.post('/generate')
async def generate_output(payload: GenerateRequest, x_session_id: str = Header(...)):
    await check_and_increment(x_session_id)

    profile_doc = db.get_relationship_profile(payload.profile_id)
    if not profile_doc:
        raise HTTPException(status_code=404, detail='Relationship Profile not found')

    kb = db.get_knowledge_base(profile_doc['supplier_id'])
    if not kb:
        raise HTTPException(status_code=404, detail='Knowledge Base not found')

    score = profile_doc.get('convergence_index', 0)
    profile = profile_doc.get('profile', {})

    if 'buying_signals' not in profile and 'buying_signals' in profile_doc:
        profile['buying_signals'] = profile_doc['buying_signals']

    if score < 60:
        raise HTTPException(status_code=400, detail={
            'message': 'Singularity Density too low to approach credibly.',
            'convergence_index': score,
            'minimum_required': 60,
        })

    try:
        if score >= 90:
            output = await generate_convergence_proposal(profile, kb)
            path, label = 'A', 'CONVERGENCE PROPOSAL'
        elif score >= 75:
            output = await generate_mutual_value_brief(profile, kb)
            path, label = 'B', 'MUTUAL VALUE BRIEF'
        else:
            output = await generate_first_signal(profile, kb)
            path, label = 'C', 'FIRST SIGNAL'
    except Exception as e:
        raise HTTPException(status_code=500,
            detail=f'Gemini generation failed: {type(e).__name__}: {e}')

    for field in ('email', 'brief', 'proposal', 'engagement_brief'):
        if output.get(field):
            output[field] = _clean(output[field])

    db.record_outcome({
        'supplier_id': profile_doc['supplier_id'],
        'profile_id': payload.profile_id,
        'company_name': profile_doc['company_name'],
        'convergence_index': score,
        'output_path': path,
        'output_label': label,
        'status': 'generated',
    })

    return {
        'path': path, 'label': label,
        'convergence_index': score,
        'company_name': profile_doc['company_name'],
        'output': output,
    }


@router.post('/export')
async def export_output(payload: ExportRequest, x_session_id: str = Header(...)):
    buying_signals = None
    score_reasoning = None
    what_would_improve = None

    if payload.include_buying_signals or payload.include_reasoning:
        profile_doc = db.get_relationship_profile(payload.profile_id)
        if profile_doc:
            profile = profile_doc.get('profile', {})
            if payload.include_buying_signals:
                signals = profile.get('buying_signals') or profile_doc.get('buying_signals')
                if signals:
                    buying_signals = signals
            if payload.include_reasoning:
                ci = profile.get('convergence_index') or {}
                if isinstance(ci, dict):
                    score_reasoning = ci.get('reasoning')
                    what_would_improve = ci.get('what_would_improve_it')

    try:
        docx_bytes = _build_docx(
            label=payload.label,
            company_name=payload.company_name,
            supplier_name=payload.supplier_name,
            convergence_index=payload.convergence_index,
            output=payload.output,
            buying_signals=buying_signals,
            score_reasoning=score_reasoning,
            what_would_improve=what_would_improve,
        )
    except RuntimeError as e:
        raise HTTPException(status_code=500, detail=str(e))

    safe = ''.join(c for c in payload.company_name if c.isalnum() or c in ' _-').strip().replace(' ', '_')
    filename = f'STRATAGENT_{safe}_{payload.label.replace(" ", "_")}.docx'
    mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type=mimetype,
        headers={'Content-Disposition': f'attachment; filename="{filename}"'},
    )


# ===========================================================================
# KB Intelligence Exports -- Visual Report, Channel Brief, Market Scan
# ===========================================================================

def _safe_filename(name: str) -> str:
    return ''.join(c for c in name if c.isalnum() or c in ' _-').strip().replace(' ', '_')


def build_visual_report_docx(kb: dict, analysis: dict) -> bytes:
    """Visual Intelligence Report -- world-class branded .docx."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
    except ImportError:
        raise RuntimeError("python-docx not installed.")

    company  = kb.get("company_name", "Supplier")
    location = kb.get("supplier_location", "")

    qi = analysis.get("quality_indicators", {})
    cp = analysis.get("competitive_position", {})
    ca = analysis.get("commercial_appeal", {})
    cf = analysis.get("channel_aesthetic_fit", {})

    tier_labels = {
        "premium":        "PREMIUM",
        "above_average":  "ABOVE AVERAGE",
        "market_average": "MARKET AVERAGE",
        "below_average":  "BELOW AVERAGE",
        "commodity":      "COMMODITY",
    }
    tier_colors = {
        "premium":        'F59E0B',
        "above_average":  '10B981',
        "market_average": '64748B',
        "below_average":  'F97316',
        "commodity":      'EF4444',
    }
    tier_key   = cp.get("tier", "market_average")
    tier_text  = tier_labels.get(tier_key, "MARKET AVERAGE")
    tier_color = tier_colors.get(tier_key, '64748B')

    doc = Document()
    _page_setup(doc)
    _add_page_footer(doc)
    _cover_block(doc, "Visual Intelligence Report", company, location)

    # Tier badge
    _section_heading(doc, "Competitive Position")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(8)
    r = p.add_run(f"  {tier_text}  ")
    r.bold = True; r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(*bytes.fromhex(tier_color))

    # Marketing description
    if analysis.get("marketing_description"):
        _section_heading(doc, "Marketing Description")
        _body_para(doc, analysis["marketing_description"], size=11)

    # Quality verdict
    if analysis.get("quality_verdict"):
        _body_para(doc, analysis["quality_verdict"], italic=True, size=10, color='475569')

    # Quality scores
    quality_rows = [
        (lbl, int(qi.get(key, 0)), 10)
        for lbl, key in [
            ("Print Clarity",    "print_clarity"),
            ("Color Richness",   "color_richness"),
            ("Composition",      "composition_score"),
            ("Production Value", "production_value"),
            ("Overall Quality",  "overall_quality"),
        ] if qi.get(key) is not None
    ]
    if quality_rows:
        _section_heading(doc, "Quality Indicators")
        _score_table(doc, quality_rows)
        if qi.get("quality_notes"):
            _body_para(doc, qi["quality_notes"], italic=True, size=9.5, color='475569')

    # Commercial appeal
    appeal_rows = [
        (lbl, int(ca.get(key, 0)), 10)
        for lbl, key in [
            ("Gift Potential",    "gift_potential"),
            ("Wall Art Appeal",   "wall_art_appeal"),
            ("Collector Appeal",  "collector_appeal"),
            ("Retail Display",    "retail_display_impact"),
        ] if ca.get(key) is not None
    ]
    if appeal_rows:
        _section_heading(doc, "Commercial Appeal")
        _score_table(doc, appeal_rows)
        if ca.get("notes"):
            _body_para(doc, ca["notes"], italic=True, size=9.5, color='475569')

    # Competitive strengths + watch points
    if cp.get("differentiators"):
        _section_heading(doc, "Competitive Strengths")
        _bullet_list(doc, cp["differentiators"], bullet='+', color='065F46')

    if cp.get("weaknesses"):
        _section_heading(doc, "Watch Points")
        _bullet_list(doc, cp["weaknesses"], bullet='!', color='9A3412')

    # Competitive benchmarks
    bench = [
        (lbl, cp.get(key))
        for key, lbl in [
            ("vs_pod_platforms",   "vs POD Platforms (Redbubble / Society6 avg)"),
            ("vs_etsy_artisan",    "vs Top Etsy Artisan Sellers"),
            ("vs_licensed_novelty","vs Licensed Novelty (Hot Topic tier)"),
        ] if cp.get(key)
    ]
    if bench:
        _section_heading(doc, "Competitive Benchmarks")
        _stats_table(doc, bench, col_widths=[2.5, 4.27])

    # Channel aesthetic fit
    if cf:
        _section_heading(doc, "Channel Aesthetic Fit")
        _channel_fit_table(doc, cf)

    # Recommended positioning
    if analysis.get("recommended_positioning"):
        _section_heading(doc, "Recommended Positioning")
        _body_para(doc, analysis["recommended_positioning"], size=11, color='0E3A58')

    _doc_footer(doc)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_channel_brief_docx(kb: dict, brief_text: str, channel_name: str = "",
                             image_b64: str = "", image_mime: str = "image/jpeg") -> bytes:
    """Channel Strategy Brief -- world-class branded .docx."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
    except ImportError:
        raise RuntimeError("python-docx not installed.")

    company  = kb.get("company_name", "Supplier")
    subtitle = f"Channel: {channel_name}" if channel_name else ""

    doc = Document()
    _page_setup(doc)
    _add_page_footer(doc)
    _cover_block(doc, "Channel Strategy Brief", company, subtitle=subtitle)

    # Product image thumbnail -- placed immediately after cover so reader sees
    # which product the brief refers to before reading any analysis.
    if image_b64:
        import base64 as _b64, tempfile as _tmp, os as _os
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        try:
            img_bytes = _b64.b64decode(image_b64)
            ext = "jpg" if "jpeg" in image_mime.lower() else image_mime.split("/")[-1]
            with _tmp.NamedTemporaryFile(suffix=f".{ext}", delete=False) as tf:
                tf.write(img_bytes)
                tmp_path = tf.name
            thumb_para = doc.add_paragraph()
            thumb_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = thumb_para.add_run()
            run.add_picture(tmp_path, width=Inches(2.4))
            _os.unlink(tmp_path)
            cap = doc.add_paragraph()
            cap.paragraph_format.space_after = Pt(12)
            cap_run = cap.add_run(kb.get("company_name", "Product"))
            cap_run.italic = True
            cap_run.font.size = Pt(8.5)
            from docx.shared import RGBColor as _RGB2
            cap_run.font.color.rgb = _RGB2(0x94, 0xA3, 0xB8)
        except Exception:
            pass  # never break the export over a missing thumbnail

    # Render brief text: detect markdown headings and bullets
    import re as _re
    blocks = brief_text.split('\n\n')
    current_section_started = False

    for block in blocks:
        block = block.strip()
        if not block:
            continue

        # Markdown heading -> section header
        # IMPORTANT: only use the first line as the heading.
        # If Gemini puts body text on the next line without a blank line,
        # block.lstrip('#') would grab the whole block. Split first.
        if block.startswith('## ') or block.startswith('# '):
            blines = block.split('\n')
            heading = blines[0].lstrip('#').strip()
            _section_heading(doc, heading)
            current_section_started = True
            # Render any lines after the heading as body text
            remainder_lines = [l.strip() for l in blines[1:] if l.strip()]
            for line in remainder_lines:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after  = Pt(6)
                parts = _re.split(r'\*\*(.+?)\*\*', line)
                for j, part in enumerate(parts):
                    if not part:
                        continue
                    r = p.add_run(part)
                    r.bold = (j % 2 == 1)
                    r.font.size = Pt(10.5)
                    r.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

        # Bullet list block
        elif _re.match(r'^[\-\*]\s', block):
            if not current_section_started:
                _section_heading(doc, "Strategic Brief")
                current_section_started = True
            lines = [l.lstrip('-* ').strip() for l in block.split('\n') if l.strip()]
            _bullet_list(doc, lines)

        else:
            if not current_section_started:
                _section_heading(doc, "Strategic Brief")
                current_section_started = True

            # Inline bold (**text**)
            for line in block.split('\n'):
                line = line.strip()
                if not line:
                    continue
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after  = Pt(6)
                parts = _re.split(r'\*\*(.+?)\*\*', line)
                for j, part in enumerate(parts):
                    if not part:
                        continue
                    r = p.add_run(part)
                    r.bold = (j % 2 == 1)
                    r.font.size = Pt(10.5)
                    r.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    _doc_footer(doc)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_market_scan_docx(kb: dict, scan: dict) -> bytes:
    """Market Intelligence Scan -- world-class branded .docx."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
    except ImportError:
        raise RuntimeError("python-docx not installed.")

    company = kb.get("company_name", "Supplier")
    product = scan.get("product", "")
    geo     = scan.get("geography_used") or scan.get("geography", "")

    doc = Document()
    _page_setup(doc)
    _add_page_footer(doc)
    _cover_block(doc, "Market Intelligence Scan", company, subtitle=product)

    # Scan summary stats
    _section_heading(doc, "Scan Summary")
    _stats_table(doc, [
        ("Product",       product),
        ("Geography",     geo),
        ("Signals Found", str(scan.get("signals_found", 0))),
        ("Open Channels", str(len(scan.get("open_channels", [])))),
    ])

    # Channel saturation
    sat = scan.get("saturation_by_channel", {})
    if sat:
        _section_heading(doc, "Channel Saturation")
        rows = [(ch, score, 100) for ch, score in sorted(sat.items(), key=lambda x: x[1])]
        _score_table(doc, rows)

    # Open channel opportunities
    open_ch = scan.get("open_channels", [])
    if open_ch:
        _section_heading(doc, "Open Channel Opportunities")
        _bullet_list(doc, open_ch, bullet='→', color='065F46')

    # Top intelligence signals
    signals = scan.get("top_signals", [])
    if signals:
        _section_heading(doc, "Top Intelligence Signals")
        for sig in signals:
            # Signal headline
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(2)
            _keep_with_next_para(p)
            sig_type = sig.get('signal_type', '')
            headline = sig.get('headline', '')
            r = p.add_run(f"[{sig_type}]  " if sig_type else "")
            r.bold = True; r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
            r2 = p.add_run(headline)
            r2.bold = True; r2.font.size = Pt(10.5)
            r2.font.color.rgb = RGBColor(0xE8, 0x7A, 0x00)

            if sig.get("detail"):
                _body_para(doc, sig["detail"], size=9.5, color='475569')

            if sig.get("channel"):
                ch_p = doc.add_paragraph()
                ch_p.paragraph_format.space_before = Pt(0)
                ch_p.paragraph_format.space_after  = Pt(8)
                r3 = ch_p.add_run(f"Channel: {sig['channel']}")
                r3.italic = True; r3.font.size = Pt(8.5)
                r3.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
            else:
                sp = doc.add_paragraph()
                sp.paragraph_format.space_after = Pt(6)

    _doc_footer(doc)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Export endpoints
# ---------------------------------------------------------------------------

class KBExportRequest(BaseModel):
    pass  # supplier_id from path; data fetched from Firestore


@router.post('/export-visual-report/{supplier_id}')
async def export_visual_report(supplier_id: str, x_session_id: str = Header(...)):
    """Export Visual Intelligence analysis as branded .docx"""
    from services import firestore as _db
    await check_and_increment(x_session_id)
    kb = _db.get_knowledge_base(supplier_id)
    if not kb:
        raise HTTPException(status_code=404, detail="Knowledge Base not found")
    analysis = kb.get("visual_analysis")
    if not analysis:
        raise HTTPException(status_code=400, detail="No visual analysis found. Run analysis first.")
    try:
        docx_bytes = build_visual_report_docx(kb, analysis)
    except RuntimeError as e:
        raise HTTPException(status_code=500, detail=str(e))
    safe = _safe_filename(kb.get("company_name", supplier_id))
    filename = f"STRATAGENT_{safe}_Visual_Intelligence.docx"
    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={'Content-Disposition': f'attachment; filename="{filename}"'},
    )


class ChannelBriefExportRequest(BaseModel):
    brief: str
    channel_name: str = ""


@router.post('/export-channel-brief/{supplier_id}')
async def export_channel_brief(supplier_id: str, payload: ChannelBriefExportRequest,
                                x_session_id: str = Header(...)):
    """Export Channel Strategy Brief as branded .docx"""
    from services import firestore as _db
    await check_and_increment(x_session_id)
    kb = _db.get_knowledge_base(supplier_id)
    if not kb:
        raise HTTPException(status_code=404, detail="Knowledge Base not found")
    # Fetch first product image for thumbnail
    img_b64 = ""
    img_mime = "image/jpeg"
    try:
        images = _db.get_product_images(supplier_id)
        if images:
            first = images[0]
            img_b64  = first.get("data", "")
            img_mime = first.get("content_type", "image/jpeg")
    except Exception:
        pass  # non-fatal -- brief still exports without thumbnail

    try:
        docx_bytes = build_channel_brief_docx(kb, payload.brief, payload.channel_name,
                                              img_b64, img_mime)
    except RuntimeError as e:
        raise HTTPException(status_code=500, detail=str(e))
    safe = _safe_filename(kb.get("company_name", supplier_id))
    ch_safe = _safe_filename(payload.channel_name or "Channel")
    filename = f"STRATAGENT_{safe}_{ch_safe}_Channel_Brief.docx"
    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={'Content-Disposition': f'attachment; filename="{filename}"'},
    )


class MarketScanExportRequest(BaseModel):
    scan: dict


@router.post('/export-market-scan/{supplier_id}')
async def export_market_scan(supplier_id: str, payload: MarketScanExportRequest,
                              x_session_id: str = Header(...)):
    """Export Market Scan intelligence report as branded .docx"""
    from services import firestore as _db
    await check_and_increment(x_session_id)
    kb = _db.get_knowledge_base(supplier_id)
    if not kb:
        raise HTTPException(status_code=404, detail="Knowledge Base not found")
    try:
        docx_bytes = build_market_scan_docx(kb, payload.scan)
    except RuntimeError as e:
        raise HTTPException(status_code=500, detail=str(e))
    safe = _safe_filename(kb.get("company_name", supplier_id))
    filename = f"STRATAGENT_{safe}_Market_Scan.docx"
    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={'Content-Disposition': f'attachment; filename="{filename}"'},
    )
