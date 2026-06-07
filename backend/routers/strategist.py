"""
STRATAGENT -- STRATEGIST Router
Cross-pipeline AI advisor. Monday Brief + Top 3 Actions.
"""
import os
import time
from datetime import datetime
from fastapi import APIRouter, Header
from fastapi.responses import FileResponse
from services import firestore as db
from services.demo_gate import check_and_increment
from agents.strategist_agent import generate_brief

router = APIRouter()

# Folder where brief docs are saved
_BRIEFS_DIR = os.path.join(os.path.dirname(__file__), "..", "..", "docs", "Briefs")


def _ensure_briefs_dir():
    os.makedirs(_BRIEFS_DIR, exist_ok=True)


def _build_brief_docx(brief: dict, generated_at: float) -> str:
    """Build a formatted .docx from the STRATEGIST brief dict. Returns file path."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    _ensure_briefs_dir()
    date_str = datetime.fromtimestamp(generated_at).strftime("%Y-%m-%d")
    filename = f"STRATAGENT_Brief_{date_str}.docx"
    filepath = os.path.join(_BRIEFS_DIR, filename)

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    ORANGE = RGBColor(0xF5, 0x9E, 0x0B)
    DARK   = RGBColor(0x1A, 0x1A, 0x2E)

    def add_heading(text, level=1):
        p = doc.add_heading(text, level=level)
        for run in p.runs:
            run.font.color.rgb = ORANGE
            run.font.bold = True
        return p

    def add_body(text, italic=False):
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(4)
        for run in p.runs:
            run.font.size = Pt(10)
            if italic:
                run.font.italic = True
        return p

    # Cover
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("STRATAGENT")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = ORANGE

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub.add_run("Monday Intelligence Brief -- " + datetime.fromtimestamp(generated_at).strftime("%A %d %B %Y"))
    run2.font.size = Pt(11)
    run2.font.color.rgb = DARK

    score = brief.get("pipeline_score", 0)
    score_p = doc.add_paragraph()
    score_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = score_p.add_run(f"Pipeline Score: {score}/100")
    sr.font.size = Pt(11)
    sr.font.bold = True

    doc.add_paragraph()

    # This week
    add_heading("This Week", 1)
    add_body(brief.get("week_headline", ""))

    # Pipeline health
    add_heading("Pipeline Health", 2)
    add_body(brief.get("pipeline_score_reasoning", ""))

    # Top 3 actions
    actions = brief.get("top_3_actions", [])
    if actions:
        add_heading("Top 3 Actions", 2)
        for a in actions:
            p = doc.add_paragraph(style="List Number")
            run = p.add_run(a.get("action", ""))
            run.font.bold = True
            run.font.size = Pt(10)
            doc.add_paragraph(
                f"Why: {a.get('why','')}  |  Module: {a.get('module','')}  |  Effort: {a.get('effort','')}",
            ).paragraph_format.left_indent = Inches(0.25)

    # Who to call
    calls = brief.get("top_calls", [])
    if calls:
        add_heading("Who to Call This Week", 2)
        for c in calls:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(f"{c.get('company','')}  (SD {c.get('ci','')} -- {c.get('urgency','')})")
            run.font.bold = True
            run.font.size = Pt(10)
            doc.add_paragraph(c.get("why_now", "")).paragraph_format.left_indent = Inches(0.25)
            if c.get("opening_line"):
                ob = doc.add_paragraph(f'"{c["opening_line"]}"')
                ob.paragraph_format.left_indent = Inches(0.25)
                for r in ob.runs:
                    r.font.italic = True
                    r.font.size = Pt(9)

    # What changed
    changed = brief.get("what_changed", [])
    if changed:
        add_heading("What Changed", 2)
        for item in changed:
            doc.add_paragraph(item, style="List Bullet")

    # STRATAGORA
    mi = brief.get("market_intelligence", {})
    if mi and mi.get("top_market_signal"):
        add_heading("STRATAGORA Market Intelligence", 2)
        sectors = ", ".join(mi.get("active_sectors", []))
        if sectors:
            add_body(f"Active sectors: {sectors}", italic=True)
        add_body(mi.get("top_market_signal", ""))
        if mi.get("stratagora_recommendation"):
            add_body(f"Recommended action: {mi['stratagora_recommendation']}")

    # KB health
    kb = brief.get("kb_health", {})
    if kb:
        add_heading("Knowledge Base Health", 2)
        table = doc.add_table(rows=2, cols=3)
        table.style = "Table Grid"
        for i, (label, key) in enumerate([("Strongest", "strongest"), ("Weakest", "weakest"), ("Fix First", "fix_first")]):
            table.cell(0, i).text = label
            table.cell(1, i).text = kb.get(key, "")

    # Watch alerts
    alerts = brief.get("watch_alerts", [])
    if alerts:
        add_heading("Watch Alerts", 2)
        for a in alerts:
            doc.add_paragraph(a, style="List Bullet")

    # Footer
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer_p.add_run("Jason L. Smith  |  Strategic Sales International ApS  |  info@strategic.dk  |  STRATAGENT -- The Intelligence Behind Agentic Sales.")
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.save(filepath)
    return filepath


@router.get("/pipeline-snapshot")
async def get_pipeline_snapshot():
    """Aggregate cross-module data. No AI calls -- fast."""
    kbs = db.list_knowledge_bases()
    profiles = db.list_all_relationship_profiles(limit=100)
    watched = db.list_all_monitored_positions(include_dismissed=False)
    outcomes = db.list_all_outcomes(limit=30)
    kb_map = {kb["id"]: kb.get("company_name", "") for kb in kbs}
    for p in profiles:
        p["supplier_name"] = kb_map.get(p.get("supplier_id", ""), p.get("supplier_id", ""))
    return {
        "kbs": kbs,
        "kbs_count": len(kbs),
        "profiles": profiles,
        "profiles_count": len(profiles),
        "watched": watched,
        "watched_count": len(watched),
        "outcomes": outcomes,
    }


@router.get("/latest-brief")
async def get_latest_brief():
    """Return the stored STRATEGIST brief from Firestore (survives restarts)."""
    data = db.get_latest_strategist_brief()
    if not data:
        return {"brief": None, "generated_at": None, "doc_path": None}
    return {
        "brief": data.get("brief"),
        "generated_at": data.get("generated_at"),
        "doc_path": data.get("doc_path", ""),
    }


@router.get("/download-brief")
async def download_latest_brief():
    """Download the latest brief as a .docx file."""
    data = db.get_latest_strategist_brief()
    if not data:
        from fastapi import HTTPException
        raise HTTPException(status_code=404, detail="No brief found. Generate one first.")
    doc_path = data.get("doc_path", "")
    if not doc_path or not os.path.exists(doc_path):
        from fastapi import HTTPException
        raise HTTPException(status_code=404, detail="Brief file not found. Regenerate the brief.")
    filename = os.path.basename(doc_path)
    return FileResponse(doc_path, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename=filename)


@router.post("/brief")
async def generate_monday_brief(x_session_id: str = Header(...)):
    """Generate the STRATEGIST Monday Brief. Saves to Firestore + docs/Briefs/ folder."""
    await check_and_increment(x_session_id)

    kbs      = db.list_knowledge_bases()
    profiles = db.list_all_relationship_profiles(limit=100)
    watched  = db.list_all_monitored_positions(include_dismissed=False)
    outcomes = db.list_all_outcomes(limit=30)
    market_signals = db.get_recent_signals_for_strategist(days=30)

    pipeline_data = {
        "kbs": kbs, "profiles": profiles,
        "watched": watched, "outcomes": outcomes,
        "market_signals": market_signals,
    }

    try:
        brief = await generate_brief(pipeline_data)
    except Exception as e:
        import traceback
        print("STRATEGIST BRIEF ERROR:\n", traceback.format_exc())
        brief = {
            "week_headline": f"Brief generation failed: {str(e)[:120]}",
            "pipeline_score": 0,
            "pipeline_score_reasoning": "Gemini API error -- retry in a few minutes.",
            "top_calls": [], "top_3_actions": [], "what_changed": [],
            "kb_health": {"strongest": "", "weakest": "", "fix_first": "Retry after 2-3 minutes."},
            "watch_alerts": [],
            "market_intelligence": {"active_sectors": [], "top_market_signal": None, "stratagora_recommendation": None},
        }

    generated_at = time.time()

    # Save to Firestore for persistence
    doc_path = ""
    try:
        doc_path = _build_brief_docx(brief, generated_at)
        print("STRATEGIST brief saved to:", doc_path)
    except Exception as e:
        print("STRATEGIST docx export failed:", e)

    db.save_strategist_brief(brief, generated_at, doc_path)

    return {
        "brief": brief,
        "generated_at": generated_at,
        "market_signals_count": len(market_signals),
        "doc_ready": bool(doc_path and os.path.exists(doc_path)),
    }
