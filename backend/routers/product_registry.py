"""
STRATAGENT -- Product Registry & Analysis Vault Router

Endpoints:
  POST   /api/product-registry/register              Register a new product
  GET    /api/product-registry/                       List all products
  GET    /api/product-registry/{id}                   Get product profile
  PATCH  /api/product-registry/{id}                   Update product fields (archetype, purpose, etc.)
  DELETE /api/product-registry/{id}                   Delete product
  POST   /api/product-registry/route                  Three-question archetype router
  POST   /api/product-registry/{id}/image             Upload image to product
  GET    /api/product-registry/{id}/images            Get product images
  DELETE /api/product-registry/{id}/images/{img_id}  Delete product image
  POST   /api/product-registry/{id}/add-url           Add URL source to product
  POST   /api/product-registry/{id}/scan              Run analysis (new vault entry)
  GET    /api/product-registry/{id}/vault             Current vault entry (locked)
  GET    /api/product-registry/{id}/vault/history     All vault versions
"""
import base64
import uuid
import time
from fastapi import APIRouter, HTTPException, UploadFile, File, Form
from fastapi.responses import StreamingResponse
import tempfile
import os
from pydantic import BaseModel, Field
from typing import Optional, List
from services import firestore as db
from agents import product_analysis_agent as paa

router = APIRouter(prefix="/api/product-registry", tags=["Product Registry"])

PURPOSE_OPTIONS = {"own_product", "affiliate_evaluation", "client_product"}


# ---------------------------------------------------------------------------
# Request models
# ---------------------------------------------------------------------------

class RegisterProductRequest(BaseModel):
    product_id: str = Field(..., description="Unique slug e.g. 'ungunk-001'")
    product_name: str
    description: str = Field(..., description="One-line description -- agents classify from this")
    archetype: Optional[str] = Field(None, description="Auto-routed if omitted")
    purpose: str = Field("own_product", description="own_product | affiliate_evaluation | client_product")
    supplier_id: Optional[str] = Field(None, description="Link to existing KB supplier")
    geography: Optional[str] = None

    # Archetype 1 -- Consumer Art / Novelty
    category: Optional[str] = None
    theme_subject: Optional[str] = None
    price_point: Optional[str] = None
    production_method: Optional[str] = None
    buyer_type: Optional[str] = None
    seasonal_relevance: Optional[str] = None

    # Archetype 2 -- Consumer Design Product
    design_story: Optional[str] = None
    sustainability_claims: Optional[str] = None
    production_location: Optional[str] = None
    b2b_potential: Optional[str] = None
    target_consumer: Optional[str] = None

    # Archetype 3 -- B2B Training / Professional
    training_subject: Optional[str] = None
    format: Optional[str] = None
    target_industries: Optional[List[str]] = None
    target_functions: Optional[List[str]] = None
    certification_relevance: Optional[str] = None
    anchor_client: Optional[str] = None
    licensing_model: Optional[str] = None
    price_point_unit: Optional[str] = None
    price_point_site: Optional[str] = None
    companion_app: Optional[str] = None

    # Archetype 4 -- B2B Industrial Supply
    kb_id: Optional[str] = None

    tags: Optional[List[str]] = None


class PatchProductRequest(BaseModel):
    archetype: Optional[str] = None
    purpose: Optional[str] = None
    description: Optional[str] = None
    geography: Optional[str] = None
    supplier_id: Optional[str] = None


class RouteRequest(BaseModel):
    keywords: List[str]
    buyer: Optional[str] = None
    product_type: Optional[str] = None
    description: Optional[str] = None


class ScanRequest(BaseModel):
    trigger_reason: str = Field("CLIENT_REQUEST")
    scan_focus: Optional[str] = Field(
        None,
        description=(
            "Free-text clarification injected into the agent prompt before scanning. "
            "Use to correct what the product actually is, narrow the search scope, "
            "or highlight specific channels/markets to prioritise."
        ),
    )


class AddUrlRequest(BaseModel):
    url: str
    label: Optional[str] = None


# ---------------------------------------------------------------------------
# Routes -- registry CRUD
# ---------------------------------------------------------------------------

@router.post("/register")
def register_product(req: RegisterProductRequest):
    # Auto-route archetype from description if not supplied
    archetype = req.archetype
    if not archetype:
        # Pass the full description as one element so multi-word tokens match
        keywords = [req.description or ""] + list(req.tags or [])
        archetype = db.route_to_archetype(keywords)
        if not archetype:
            # Keyword router found no match -- use sync Gemini fallback
            archetype = paa.classify_archetype_with_gemini_sync(
                req.description or "", req.product_name
            )
        if not archetype:
            archetype = "consumer_design_product"  # safe default

    if archetype not in db.VALID_ARCHETYPES:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid archetype '{archetype}'. Must be one of: {sorted(db.VALID_ARCHETYPES)}"
        )
    if req.purpose not in PURPOSE_OPTIONS:
        raise HTTPException(status_code=400, detail=f"purpose must be one of {PURPOSE_OPTIONS}")

    data = req.dict(exclude_none=True)
    data["archetype"] = archetype
    data["archetype_label"] = paa.ARCHETYPE_LABELS.get(archetype, archetype)
    db.save_product_registry(req.product_id, data)
    return {
        "status": "registered",
        "product_id": req.product_id,
        "archetype": archetype,
        "archetype_label": data["archetype_label"],
    }


@router.get("/")
def list_products():
    products = db.list_product_registry()
    return {"products": products, "count": len(products)}


@router.get("/{product_id}")
def get_product(product_id: str):
    product = db.get_product_registry(product_id)
    if not product:
        raise HTTPException(status_code=404, detail=f"Product not found: {product_id}")
    return product


@router.patch("/{product_id}")
def patch_product(product_id: str, req: PatchProductRequest):
    """Update mutable fields on an existing product. Use to fix archetype after registration."""
    product = db.get_product_registry(product_id)
    if not product:
        raise HTTPException(status_code=404, detail=f"Product not found: {product_id}")
    updates = req.dict(exclude_none=True)
    if "archetype" in updates:
        if updates["archetype"] not in db.VALID_ARCHETYPES:
            raise HTTPException(
                status_code=400,
                detail=f"Invalid archetype. Must be one of: {sorted(db.VALID_ARCHETYPES)}"
            )
        updates["archetype_label"] = paa.ARCHETYPE_LABELS.get(updates["archetype"], updates["archetype"])
    if "purpose" in updates and updates["purpose"] not in PURPOSE_OPTIONS:
        raise HTTPException(status_code=400, detail=f"purpose must be one of {PURPOSE_OPTIONS}")
    db.save_product_registry(product_id, updates)
    return {"status": "updated", "product_id": product_id, "updated_fields": list(updates.keys())}


@router.delete("/{product_id}")
def delete_product(product_id: str):
    if not db.get_product_registry(product_id):
        raise HTTPException(status_code=404, detail=f"Product not found: {product_id}")
    db.delete_product_registry(product_id)
    return {"status": "deleted", "product_id": product_id}


# ---------------------------------------------------------------------------
# Archetype router
# ---------------------------------------------------------------------------

@router.post("/route")
def route_archetype(req: RouteRequest):
    archetype = None
    if req.buyer and req.product_type:
        buyer, ptype = req.buyer.lower(), req.product_type.lower()
        if buyer == "consumer":
            archetype = "consumer_art_novelty" if ptype == "art" else "consumer_design_product"
        elif buyer == "business":
            archetype = "b2b_training_professional" if ptype == "training" else "b2b_industrial_supply"

    if not archetype:
        # Pass description as full string so multi-word tokens match
        kw = [req.description or ""] + list(req.keywords)
        archetype = db.route_to_archetype(kw)

    if not archetype and req.description:
        # Sync Gemini fallback for novel/prototype products
        archetype = paa.classify_archetype_with_gemini_sync(req.description)

    if not archetype:
        archetype = "consumer_design_product"

    return {
        "archetype": archetype,
        "archetype_label": paa.ARCHETYPE_LABELS.get(archetype, archetype),
        "required_intake_fields": _required_fields_for_archetype(archetype),
        "signal_types": paa.ARCHETYPE_SIGNALS.get(archetype, []),
    }


# ---------------------------------------------------------------------------
# Image upload / management
# ---------------------------------------------------------------------------

@router.post("/{product_id}/image")
async def upload_product_image(
    product_id: str,
    file: UploadFile = File(...),
    label: str = Form(""),
    tags: str = Form(""),
):
    product = db.get_product_registry(product_id)
    if not product:
        raise HTTPException(status_code=404, detail=f"Product not found: {product_id}")

    content = await file.read()
    if len(content) > 5 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="Image too large (max 5 MB)")

    # Auto-resize if needed
    try:
        from PIL import Image as PILImage
        import io
        if len(content) > 700_000:
            img = PILImage.open(io.BytesIO(content))
            img.thumbnail((1600, 1600), PILImage.LANCZOS)
            buf = io.BytesIO()
            fmt = img.format or "JPEG"
            img.save(buf, format=fmt, quality=85)
            content = buf.getvalue()
    except Exception:
        pass

    image_id = str(uuid.uuid4())
    db.save_product_image(image_id, {
        "product_id": product_id,
        "product_name": product.get("product_name", ""),
        "supplier_id": product.get("supplier_id"),
        "label": label,
        "tags": tags,
        "filename": file.filename,
        "content_type": file.content_type,
        "data": base64.b64encode(content).decode("utf-8"),
    })
    return {"status": "uploaded", "image_id": image_id, "product_id": product_id}


@router.get("/{product_id}/images")
def get_product_images(product_id: str):
    images = db.get_product_registry_images(product_id)
    # Strip base64 data from list view for performance
    slim = [{k: v for k, v in img.items() if k != "data"} for img in images]
    return {"images": slim, "count": len(slim)}


@router.delete("/{product_id}/images/{image_id}")
def delete_product_image(product_id: str, image_id: str):
    db.delete_product_image(image_id)
    return {"status": "deleted", "image_id": image_id}


# ---------------------------------------------------------------------------
# URL source ingestion
# ---------------------------------------------------------------------------

@router.post("/{product_id}/add-url")
def add_product_url(product_id: str, req: AddUrlRequest):
    product = db.get_product_registry(product_id)
    if not product:
        raise HTTPException(status_code=404, detail=f"Product not found: {product_id}")
    urls = product.get("source_urls", [])
    urls.append({"url": req.url, "label": req.label or req.url, "added_at": time.time()})
    db.save_product_registry(product_id, {"source_urls": urls})
    return {"status": "added", "url": req.url, "product_id": product_id}


# ---------------------------------------------------------------------------
# Analysis vault
# ---------------------------------------------------------------------------

@router.post("/{product_id}/scan")
async def scan_product(product_id: str, req: ScanRequest):
    if req.trigger_reason not in db.VALID_TRIGGER_REASONS:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid trigger_reason. Must be one of: {sorted(db.VALID_TRIGGER_REASONS)}"
        )
    try:
        vault_entry = await paa.run_product_analysis(
            product_id, req.trigger_reason, scan_focus=req.scan_focus
        )
        return {
            "status": "analysis_complete",
            "vault_entry_id": vault_entry.get("id"),
            "version": vault_entry.get("version"),
            "signal_count": vault_entry.get("analysis", {}).get("signal_count", 0),
            "trigger_reason": req.trigger_reason,
        }
    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Scan failed: {str(e)}")


@router.get("/{product_id}/vault")
def get_vault(product_id: str):
    entry = db.get_current_vault_entry(product_id)
    if not entry:
        raise HTTPException(
            status_code=404,
            detail=f"No analysis found for '{product_id}'. Run /scan first."
        )
    return entry


@router.get("/{product_id}/vault/history")
def get_vault_history(product_id: str):
    history = db.get_vault_history(product_id)
    return {"product_id": product_id, "versions": history, "count": len(history)}



@router.get("/{product_id}/vault/buyer-targets")
async def get_buyer_targets(product_id: str):
    """
    Extract buyer targets from the current vault analysis.
    Returns aggregated sectors, geography, and suggested FI prospect names.
    """
    entry = db.get_current_vault_entry(product_id)
    if not entry:
        raise HTTPException(status_code=404, detail="No analysis found. Run /scan first.")

    product = db.get_product_registry(product_id) or {}
    analysis = entry.get("analysis", {})
    signals = analysis.get("signals", [])

    targets = await paa.extract_buyer_targets(signals, product)
    return {
        "product_id": product_id,
        "product_name": analysis.get("product_name", product_id),
        **targets,
    }


@router.get("/{product_id}/vault/export-docx")
def export_vault_docx(product_id: str):
    """Export the current vault analysis as a STRATAGENT-branded Product Market Brief."""
    entry = db.get_current_vault_entry(product_id)
    if not entry:
        raise HTTPException(status_code=404, detail="No analysis found. Run /scan first.")

    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from datetime import datetime as dt

    analysis   = entry.get("analysis", {})
    signals    = analysis.get("signals", [])
    archetype  = analysis.get("archetype_label", analysis.get("archetype", ""))
    prod_name  = analysis.get("product_name", product_id)
    scan_ts    = analysis.get("scan_timestamp", "")
    version    = entry.get("version", 1)
    trigger    = entry.get("trigger_reason", "manual")
    scan_focus = analysis.get("scan_focus") or ""

    # -- Brand colours (STRATAGENT / PERSOLIT standard) ----------------------
    ORANGE     = RGBColor(0xE8, 0x7A, 0x00)   # primary accent
    ORANGE_DK  = RGBColor(0xB2, 0x3B, 0x00)   # urgency / strength tag
    BLACK      = RGBColor(0x00, 0x00, 0x00)
    GRAY       = RGBColor(0x6B, 0x6B, 0x6B)   # metadata / timestamps
    GRAY_LIGHT = RGBColor(0x9A, 0x9A, 0x9A)   # footer

    RULE_CHAR  = "-" * 60  # horizontal rule made of box-drawing dashes

    doc = Document()

    # -- Page setup (A4, 0.76" margins) --------------------------------------
    for sec in doc.sections:
        sec.page_height   = Inches(11.69)
        sec.page_width    = Inches(8.27)
        sec.top_margin    = Inches(0.76)
        sec.bottom_margin = Inches(0.76)
        sec.left_margin   = Inches(0.90)
        sec.right_margin  = Inches(0.90)

    # -- Helpers -------------------------------------------------------------
    def _r(para, text, bold=False, size=10, color=None, italic=False):
        r = para.add_run(text)
        r.font.name   = "Arial"
        r.font.size   = Pt(size)
        r.font.bold   = bold
        r.font.italic = italic
        r.font.color.rgb = color if color else BLACK
        return r

    def _para(text="", bold=False, size=10, color=None,
              align=WD_ALIGN_PARAGRAPH.LEFT, sb=0, sa=3):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(sb)
        p.paragraph_format.space_after  = Pt(sa)
        if text:
            _r(p, text, bold=bold, size=size, color=color)
        return p

    def _rule(sb=4, sa=4):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(sb)
        p.paragraph_format.space_after  = Pt(sa)
        _r(p, RULE_CHAR, size=8, color=RGBColor(0xCC, 0xCC, 0xCC))

    def _section(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(4)
        _r(p, text, bold=True, size=11, color=ORANGE)

    # -- Metadata table ------------------------------------------------------
    scan_date = scan_ts[:10] if scan_ts else dt.utcnow().strftime("%Y-%m-%d")
    tbl = doc.add_table(rows=5, cols=2)
    tbl.style = "Table Grid"
    meta_rows = [
        ("PRODUCT",      prod_name),
        ("ARCHETYPE",    archetype),
        ("SCAN DATE",    scan_date),
        ("VERSION",      "v{} -- {}".format(version, trigger.replace("_", " ").title())),
        ("PREPARED BY",  "Jason L. Smith  |  Strategic Sales International ApS"),
    ]
    for i, (lbl, val) in enumerate(meta_rows):
        row = tbl.rows[i]
        lr = row.cells[0].paragraphs[0].add_run(lbl)
        lr.font.name = "Arial"; lr.font.size = Pt(9); lr.font.bold = True
        lr.font.color.rgb = ORANGE
        vr = row.cells[1].paragraphs[0].add_run(val)
        vr.font.name = "Arial"; vr.font.size = Pt(9)
    _para("")

    # -- Title ---------------------------------------------------------------
    _rule(sb=0, sa=6)
    _para("PRODUCT MARKET BRIEF", bold=True, size=22, color=ORANGE,
          align=WD_ALIGN_PARAGRAPH.LEFT, sb=0, sa=6)
    _rule(sb=0, sa=8)

    # -- Executive Summary ---------------------------------------------------
    _section("EXECUTIVE SUMMARY")
    n = len(signals)
    types_found = list(dict.fromkeys(
        s.get("signal_type", "") for s in signals if s.get("signal_type")
    ))
    _para(
        "{} market signal{} detected across {} signal type{}. "
        "Types: {}{}.".format(
            n, "s" if n != 1 else "",
            len(types_found), "s" if len(types_found) != 1 else "",
            ", ".join(types_found[:6]),
            "..." if len(types_found) > 6 else ""
        ),
        size=10, sb=2, sa=6
    )
    if scan_focus:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(6)
        _r(p, "Scan focus: ", bold=True, size=9, color=GRAY)
        _r(p, scan_focus, bold=False, size=9, color=GRAY, italic=True)

    # -- Signals -------------------------------------------------------------
    _rule()
    _section("BUYING SIGNALS DETECTED  ({} signals)".format(n))

    for i, sig in enumerate(signals, 1):
        sig_type   = sig.get("signal_type", "SIGNAL")
        headline   = sig.get("headline", sig.get("title", sig.get("summary", "")))
        detail     = sig.get("detail", "")
        action     = sig.get("action", "")
        source     = sig.get("source_url", sig.get("source", ""))
        urgency    = sig.get("urgency", sig.get("confidence", ""))
        channel    = sig.get("channel", sig.get("industry_vertical", "general"))

        # Signal type header line
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(2)
        _r(p, sig_type, bold=True, size=10, color=ORANGE)
        if urgency:
            _r(p, "  {} STRENGTH".format(urgency), bold=True, size=8, color=ORANGE_DK)
        if channel and channel.lower() != "general":
            _r(p, "  -- {}".format(channel), bold=False, size=9, color=GRAY)

        if headline:
            _para(headline, bold=True, size=10, sb=0, sa=2)

        if detail:
            _para(detail, size=10, sb=0, sa=3)

        if action:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(2)
            _r(p, "Next step: ", bold=True, size=9, color=ORANGE)
            _r(p, action, bold=False, size=9)

        if source:
            _para("Source: {}".format(source[:120]), size=8, color=GRAY, sb=0, sa=2)

        if i < n:
            _rule(sb=6, sa=2)

    # -- Footer --------------------------------------------------------------
    _rule(sb=10, sa=6)
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.LEFT
    foot.paragraph_format.space_before = Pt(0)
    foot.paragraph_format.space_after  = Pt(0)
    _r(foot, "Jason L. Smith  |  Strategic Sales International ApS\n"
             "info@strategic.dk  |  www.strategic.dk  |  +45 24 99 23 93\n"
             "CVR: 41945621  |  4000 Roskilde, Denmark\n"
             "STRATAGENT -- The Intelligence Behind Agentic Sales",
       size=8, color=GRAY_LIGHT)

    # -- Save & stream -------------------------------------------------------
    safe = prod_name.replace(" ", "_").lower()
    filename = "{}_{}_market_brief_v{}.docx".format(safe, scan_date, version)

    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    doc.save(tmp.name)
    tmp.close()

    def iterfile():
        with open(tmp.name, "rb") as f:
            yield from f
        os.unlink(tmp.name)

    return StreamingResponse(
        iterfile(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": 'attachment; filename="{}"'.format(filename)},
    )


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

_ARCHETYPE_INTAKE = {
    "consumer_art_novelty": [
        "product_name", "category", "theme_subject", "price_point",
        "production_method", "buyer_type", "seasonal_relevance", "geography",
    ],
    "consumer_design_product": [
        "product_name", "category", "design_story", "sustainability_claims",
        "price_point", "production_location", "b2b_potential", "target_consumer", "geography",
    ],
    "b2b_training_professional": [
        "product_name", "training_subject", "format", "target_industries",
        "target_functions", "certification_relevance", "anchor_client",
        "licensing_model", "price_point_unit", "price_point_site",
        "companion_app", "geography",
    ],
    "b2b_industrial_supply": [
        "product_name", "description", "geography",
    ],
}


def _required_fields_for_archetype(archetype: str):
    return _ARCHETYPE_INTAKE.get(archetype, ["product_name", "description"])
